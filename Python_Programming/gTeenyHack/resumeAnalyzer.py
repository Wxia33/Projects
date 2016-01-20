from docx import Document
import re
import os
from Tkinter import *
#from PIL import *


################################################################

def resumeAnalyze(directory,listKeyword): #look through direcory 
    filesDir = os.listdir(directory)
    for i in range(len(filesDir)):
        fileName = filesDir[i]
        if fileName.find('.docx') == -1:
            n = 0
        elif fileName.find('~$') == -1:
            resumeloc = directory + '\\%s' % (filesDir[i])
            document = Document(resumeloc)
            docText = '\n\n'.join([paragraph.text.encode('utf-8') for paragraph in document.paragraphs]) #outputs a string of text data extracted from word file
            loweredTxt = docText.lower()
            #print resumeloc
            employability(listKeyword, docText, loweredTxt)
    return docText, loweredTxt, resumeloc

def contactinfo(resTxt): #find mobile, email, name, address.  Not essential ATM
    mobil = int(resTxt.find('+1'))    
    mobilenumber = resTxt[mobil:mobil+15]
    return mobilenumber

def employability(listKeys, origresTxt, lowresTxt): #Finds keyword match in resume and
    #outputs name of individual and match with job keyword list
    '''
    for keys in listKeys:
        print int(loweredTxt.find(keys))
    '''
    count = 0
    for i in range(0,len(listKeys)):
        if listKeys[i] in lowresTxt:
            count += 1

    newline = [m.start() for m in re.finditer(r'\n',origresTxt)]
    name = origresTxt[0:newline[1]-1]
    nameFi = name.replace(' ','_')

    percentage = float(count) / float(len(listKeys)) * 100
    sentence = name + ' has a resume match of %d percent' % percentage
    print sentence
    return percentage

def inputKeys(listKeyword): #replace this function with gui input
    while wordkey != '@Done':
        wordkey = raw_input('Keywords to find: ')
        listKeyword.append(wordkey)
    del listKeyword[-1]
    return listKeyword
    
'''
files = os.listdir("C:\Users\Oracle\Desktop\\")
resumeloc = "C:\Users\Oracle\Desktop%s" % ('\XiaResume.docx')
document = Document(resumeloc)
docText = '\n\n'.join([paragraph.text.encode('utf-8') for paragraph in document.paragraphs]) #outputs a string of text data extracted from word file
loweredTxt = docText.lower()

listKeyword = ['research',
               'experience',
               'leadership',
               'awards',
               'organizations',
               'prototype',
               'ap scholar',
               'national merit',
               'software',
               'python',
               'design']
'''
#perc = employability(listKeyword)
info = []
def update(event):
    if(entry.get() != ""):
        for x in range(0, len(entry.get().split(", "))):
            info.append(entry.get().split(", ")[x])
            text.insert(END, entry.get().split(", ")[x] + '\n')
            entry.delete(0, 100)
    print info

def getArray():
    listKeyword = info
    docText,loweredTxt,resumeloc = resumeAnalyze('C:\Users\Oracle\Desktop',listKeyword)

root = Tk()
root.title("Resume searcher")
root.geometry("500x500");
window = Frame(root);
window.grid()

#image = root.PhotoImage("RTdRnae.gif")
#background_label = tk.Label(root, image=image)
#background_label.place(x=0, y=0, relwidth=500, relheight=500)

button = Button(root, text = "Search", font = ('Arial', 15), command = getArray)
button.grid()

entry = Entry(root)
entry.grid(row = 0, column = 1)

lbl = Label(root, text = "        Write tags here", padx=20, font=('Arial', 20))
lbl.grid(row=0, column=0, columnspan = 7, sticky = W)


res = Label(root, padx=20, font = ('Arial', 20), height = 2, width = 10)
res.grid(row=1, column=2, columnspan = 1, sticky = W)

entry.bind("<Return>", update)

text = Text(root, height=10, width=200, wrap = WORD)
text.grid(row=2,column=0, columnspan = 8, sticky = W)
text.insert(END, "Tags:" + "\n")

root.mainloop()
